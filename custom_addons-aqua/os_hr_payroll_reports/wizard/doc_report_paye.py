# -*- coding: utf-8 -*-
# Copyright (C) 2017-Today  Technaureus Info Solutions(<http://technaureus.com/>).

from odoo import fields, models, api
from datetime import date,datetime
import xlsxwriter
from io import BytesIO
import base64
from odoo.tools import misc
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class ReportXlsx(models.TransientModel):
    _inherit = 'report.xlsx.wizard'
    
    report_type = fields.Selection([('epf','EPF Contribution'),('etf','ETF Contribution'),
                                    ('etf2','ETF Form II'),('payroll','Payroll Report'),
                                    ('paye1','PAYE 94 PART 1'),('paye2','PAYE 94 PART 2'),
                                    ('paye3','PAYE T-9A'),('paye4','PAYE T-10')], default='paye4')
    
    emp_id = fields.Many2one('hr.employee', string='Employee')
    
    def print_xlsx_reports(self,fl):
        if self.report_type == 'paye4':
            fl = self.print_doc_paye()
        return super(ReportXlsx, self).print_xlsx_reports(fl)

    def print_doc_paye(self):
        year_start = self.year_start
        year_end = self.year_end
        previous_yr = (datetime.strptime(str(year_start),'%Y-%m-%d').date()).strftime('%Y')
        current_yr = (datetime.strptime(str(year_end),'%Y-%m-%d').date()).strftime('%Y')
        flname = os.path.join(os.path.dirname(__file__), 'PAYE T 10 Certificate of '+previous_yr+'-'+current_yr+'('+str(datetime.today())+')'+'.docx')
        domain = [('date_from','>=',self.year_start),('date_to','<=',self.year_end)]
        domain.append(('state','=','done'))
        domain.append(('employee_id','=',self.emp_id.id))
        payslips = self.env['hr.payslip'].search(domain)
        tot_gross = tot_tax = non_cash = 0
        for payslip in payslips:
            for line in payslip.line_ids:
                if payslip.credit_note == True:
                    if line.code == 'GROSS':
                        tot_gross -= line.total
                    elif line.code in ['PAYE','PTAX']:
                        tot_tax -= line.total
                    elif line.code == 'NCA':
                        non_cash -= line.total
                else:
                    if line.code == 'GROSS':
                        tot_gross += line.total
                    elif line.code in ['PAYE','PTAX']:
                        tot_tax += line.total
                    elif line.code == 'NCA':
                        non_cash += line.total
        
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.4110236)
            section.right_margin = Inches(0.4110236)
            
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
#         font.bold = True
        p1 = document.add_paragraph('PRESCRIBED UNDER SECTION 120 OF THE INLAND REVENUE ACT.      P.A.Y.E. / T. 10 (New)')
        p1.style = document.styles['Normal']
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = document.add_paragraph('S  R  I     L  A  N  K  A     I  N  L  A  N  D     R  E  V  E  N  U  E')
        p2.style = document.styles['Normal']
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p3 = document.add_paragraph('P    .    A    .    Y    .    E')
        p3.style = document.styles['Normal']
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        p4 = document.add_paragraph('C E R T I F I C A T E     O F     I N C O M E     T A X     D E D U C T I O N S')
        p4.style = document.styles['Normal']
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p5 = document.add_paragraph('\n\nFULL NAME OF EMPLOYEE \t')
        p5.style = document.styles['Normal']
        p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p5.add_run('\t\t: \t')
        name = self.emp_id.name
        if self.emp_id.surname:
            name += " "+self.emp_id.surname
        p5.add_run(name)
        
        p6 = document.add_paragraph('EMPLOYEE NIC NO \t')
        p6.style = document.styles['Normal']
        p6.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p6.add_run('\t\t\t:\t')
        p6.add_run(self.emp_id.identification_id)
        
        p7 = document.add_paragraph("SERIAL NO OF EMPLOYEE'S PAY SHEET \t:\t")
        p7.style = document.styles['Normal']
        p7.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p7.add_run(self.emp_id.contract_id.member_no)
        p7.add_run("\tEMPLOYER'S REGISTRATION NO : ")
        p7.add_run()
        
        p8 = document.add_paragraph("PERIOD OF SERVICE FOR WHICH\nREMUNERATION WAS PAID")
        p8.style = document.styles['Normal']
        p8.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p8.add_run('\t\t\tFrom :\t')
        p8.add_run(self.year_start)
        p8.add_run('\t\tTo :\t')
        p8.add_run(self.year_end)
        
        p9 = document.add_paragraph("TOTAL AMOUNT OF GROSS REMUNERATION\nAS PER PAY SHEET")
        p9.style = document.styles['Normal']
        p9.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p9.add_run('\t\t\t\t:   Rs ')
        p9.add_run(str(tot_gross)+"0")
        
        p10 = document.add_paragraph("CASH  : Rs ")
        p10.style = document.styles['Normal']
        p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cash = tot_gross-non_cash
        p10.add_run(str(cash)+"0")
        p10.add_run("\t\t\tNON CASH  : Rs ")
        p10.add_run(str(non_cash)+"0")
        p10.add_run("\t\tBENEFITS : Rs ")
        p10.add_run(str(tot_gross)+"0")
        
        
        p11 = document.add_paragraph("TOTAL AMOUNT TAX DEDUCTED")
        p11.style = document.styles['Normal']
        p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p11.add_run('\t\t: Rs ')
        p11.add_run(str(tot_tax)+"0 ")
        p11.add_run('ONLY')
        
        p12 = document.add_paragraph("TOTAL AMOUNT TAX REMITTED TO THE\nDEPT. OF INLAND REVENUE")
        p12.style = document.styles['Normal']
        p12.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p12.add_run('\t\t\t: Rs ')
        p12.add_run(str(tot_tax)+"0 ")
        p12.add_run('ONLY')
        
        p13 = document.add_paragraph("\n\nI certify the above particulars as correct")
        p13.style = document.styles['Normal']
        p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p13.paragraph_format.left_indent = Inches(2.50)
        
        p14 = document.add_paragraph("Signature of Employer : ")
        p14.style = document.styles['Normal']
        p14.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p14.paragraph_format.left_indent = Inches(1.65)
        
        p15 = document.add_paragraph("Name of Employer   : ")
        p15.style = document.styles['Normal']
        p15.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p15.add_run(self.env.user.company_id.name)
        p15.paragraph_format.left_indent = Inches(2.40)
        
        p16 = document.add_paragraph("Address  : ")
        p16.style = document.styles['Normal']
        p16.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p16.paragraph_format.left_indent = Inches(1.95)
        p16.add_run(str(self.env.user.company_id.street))
        
        if self.env.user.company_id.street2:
            p16.add_run(str(", "+self.env.user.company_id.street2))
            
        if self.env.user.company_id.city:
            p17 = document.add_paragraph("")
            p17.style = document.styles['Normal']
            p17.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p17.paragraph_format.left_indent = Inches(3.15)
            p17.add_run(str(self.env.user.company_id.city))
        if self.env.user.company_id.state_id:
            p17.add_run(str(", "+self.env.user.company_id.state_id.name))
        if self.env.user.company_id.zip:
            p17.add_run(str("-"+self.env.user.company_id.zip))
            
        if self.env.user.company_id.country_id:
            p18 = document.add_paragraph("")
            p18.style = document.styles['Normal']
            p18.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p18.paragraph_format.left_indent = Inches(2.05)
            p18.add_run(str(self.env.user.company_id.country_id.name))
                
        p19 = document.add_paragraph("Date  :\t")
        p19.style = document.styles['Normal']
        p19.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p19.add_run(str(date.today()))
        
        p20 = document.add_paragraph("N.B.   - ")
        p20.style = document.styles['Normal']
        p20.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p20.add_run('1. THIS CERTIFICATE SERVES AS PROOF OF PAYMENT')
        
        p21 = document.add_paragraph("\t2. ATTACH THIS CERTIFICATE (I) TO YOUR RETURN FOR THE RELEVANT YEAR OF\n\t ASSESMENT IF YOU ARE ")
        p21.style = document.styles['Normal']
        p21.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p21.add_run('REQUIRED TO FURNISH A RETURN , OR (II) TO CLAIM REFUND, IF ANY')
        
#         document.add_page_break()
        
        document.save(flname)
        return flname
